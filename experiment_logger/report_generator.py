from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from aliases import ALIASES


class ReportGenerator:

    def __init__(self, experiment, output_folder="reports"):

        self.experiment = experiment

        # Correct
        self.output_folder = Path(output_folder)

        # Create reports folder automatically
        self.output_folder.mkdir(parents=True, exist_ok=True)

    ############################################################

    def create_table(self, document, title, data):

        # Skip empty sections
        if not data:
            return

        document.add_heading(title, level=2)

        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        header = table.rows[0].cells
        header[0].text = "Parameter"
        header[1].text = "Value"

        for key, value in data.items():

            # Skip empty values
            if value in [None, "", "N/A"]:
                continue

            row = table.add_row().cells

            display_name = ALIASES.get(
                key,
                key.replace("_", " ").title()
            )

            row[0].text = display_name
            row[1].text = str(value)

        document.add_paragraph()

    ############################################################

    def generate(self):

        document = Document()

        # Font
        document.styles["Normal"].font.name = "Calibri"
        document.styles["Normal"].font.size = Pt(11)

        # Title
        heading = document.add_heading(
            "AI Training Report",
            level=1
        )
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Report generation date
        document.add_paragraph(
            f"Generated On: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}"
        )

        document.add_paragraph()

        ########################################################

        self.create_table(
            document,
            "Model Information",
            self.experiment.model
        )

        self.create_table(
            document,
            "Dataset Information",
            self.experiment.dataset
        )

        self.create_table(
            document,
            "Training Configuration",
            self.experiment.training
        )

        self.create_table(
            document,
            "Performance Metrics",
            self.experiment.metrics
        )

        self.create_table(
            document,
            "Additional Information",
            self.experiment.extra
        )

        ########################################################
        # Create report filename
        ########################################################

        model_name = self.experiment.model.get(
            "model_name",
            "Unknown_Model"
        )

        # Make filename safe
        model_name = "".join(
            c if c.isalnum() or c in ("_", "-") else "_"
            for c in model_name
        )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        report_name = (
            f"{model_name}_Training_Report_{timestamp}.docx"
        )

        report_path = self.output_folder / report_name

        document.save(report_path)

        return report_path
